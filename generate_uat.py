import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# ─── COLOR PALETTE DESIGN ──────────────────────────────────────────────────
COLOR_PRIMARY_HEX = "1A365D"    # Deep Academic Navy Blue
COLOR_ACCENT_HEX  = "2B6CB0"    # Crisp Threat Intel Blue
COLOR_ZEBRA_HEX   = "F7FAFC"    # Ultra-soft background tint for odd rows
COLOR_BORDER_HEX  = "CBD5E0"    # Clean Slate Border line
COLOR_TEXT_HEX    = "2D3748"    # Premium Charcoal Text color (softer than pure black)

COLOR_PRIMARY = RGBColor(0x1A, 0x36, 0x5D)
COLOR_ACCENT  = RGBColor(0x2B, 0x6C, 0xB0)
COLOR_TEXT    = RGBColor(0x2D, 0x37, 0x48)
COLOR_MUTED   = RGBColor(0x71, 0x80, 0x96)
COLOR_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)

def apply_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def apply_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def apply_table_borders(table, color_hex):
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

def add_section_banner(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(7.0)
    
    cell = table.cell(0, 0)
    apply_cell_shading(cell, COLOR_PRIMARY_HEX)
    apply_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    p_cell = cell.paragraphs[0]
    p_cell.paragraph_format.space_before = Pt(0)
    p_cell.paragraph_format.space_after = Pt(0)
    run = p_cell.add_run(text.upper())
    run.font.name = 'Segoe UI'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_WHITE

# ─── INITIALIZE MASTER DOCUMENT ──────────────────────────────────────────────
doc = Document()

# Configure Premium Wide Layout Margins
for section in doc.sections:
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin   = Inches(0.75)
    section.right_margin  = Inches(0.75)

# Set global document style defaults
style_normal = doc.styles['Normal']
style_normal.font.name = 'Segoe UI'
style_normal.font.size = Pt(10.5)
style_normal.font.color.rgb = COLOR_TEXT

# ─── HEADER TITLE BLOCK ──────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(2)

run_main = title_p.add_run("USER ACCEPTANCE TESTING (UAT) EVALUATION\n")
run_main.font.size = Pt(18)
run_main.font.bold = True
run_main.font.color.rgb = COLOR_PRIMARY

run_sub = title_p.add_run("Automated IP Risk Profiler System  |  Capstone Project")
run_sub.font.size = Pt(12)
run_sub.font.bold = True
run_sub.font.color.rgb = COLOR_ACCENT

# Thin geometric horizontal rules divider
p_hr = doc.add_paragraph()
p_hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_hr.paragraph_format.space_after = Pt(14)
hr_run = p_hr.add_run("═" * 62)
hr_run.font.color.rgb = COLOR_MUTED
hr_run.font.size = Pt(10)

# ─── RESPONDENT INFORMATION CARD ─────────────────────────────────────────────
info_table = doc.add_table(rows=2, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_table.autofit = False
info_table.columns[0].width = Inches(3.5)
info_table.columns[1].width = Inches(3.5)

fields = [
    ("Name (Optional):", "___________________________"),
    ("Role / Department:", "___________________________"),
    ("Date of Evaluation:", "___________________________"),
    ("UAT Session No.:", "___________________________")
]

idx = 0
for row in info_table.rows:
    for cell in row.cells:
        apply_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(4)
        lbl, line = fields[idx]
        
        r_lbl = p.add_run(f"{lbl} ")
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(10)
        r_lbl.font.color.rgb = COLOR_PRIMARY
        
        r_line = p.add_run(line)
        r_line.font.color.rgb = COLOR_MUTED
        idx += 1

# Instructions callout panel box matrix
p_inst = doc.add_paragraph()
p_inst.paragraph_format.space_before = Pt(18)
p_inst.paragraph_format.space_after = Pt(16)

inst_box = doc.add_table(rows=1, cols=1)
inst_box.alignment = WD_TABLE_ALIGNMENT.CENTER
inst_box.autofit = False
inst_box.columns[0].width = Inches(7.0)
cell_inst = inst_box.cell(0, 0)
apply_cell_shading(cell_inst, "F8FAFC")
apply_cell_margins(cell_inst, top=120, bottom=120, left=160, right=160)
# Left thick border line wrapper
tcPr = cell_inst._tc.get_or_add_tcPr()
tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="{COLOR_PRIMARY_HEX}"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
tcPr.append(tcBorders)

p_in_inst = cell_inst.paragraphs[0]
p_in_inst.paragraph_format.space_after = Pt(0)
r_ib = p_in_inst.add_run("EVALUATION INSTRUCTIONS:\n")
r_ib.font.bold = True
r_ib.font.size = Pt(9.5)
r_ib.font.color.rgb = COLOR_PRIMARY
r_it = p_in_inst.add_run("Please interact with the system interface deployment. For each objective parameter statement below, check or circle the matrix number that matches your assessment value where: 1 = Strongly Disagree (SD), 2 = Disagree (D), 3 = Neutral (N), 4 = Agree (A), and 5 = Strongly Agree (SA).")
r_it.font.size = Pt(9.5)
r_it.font.color.rgb = COLOR_TEXT

# ─── CORE SYSTEM QUESTIONS SCHEMA DATA MATRICES ─────────────────────────────
questions = [
    ("Ease of Use", "The system interface deployment layout is intuitive and straightforward to navigate."),
    ("Ease of Use", "I was able to execute dynamic scanning and network asset logs without external assistance."),
    ("Alert Clarity", "The raised security alert payloads communicate the exact nature and score severity clearly."),
    ("Alert Clarity", "Alert risk boundaries (Low/Medium/High) map accurately to observed asset priorities."),
    ("Dashboard Readability", "The central dashboard structure aggregates live risk feeds into a highly readable layout."),
    ("Dashboard Readability", "Interactive tracking charts and system visualization assets effectively support risk choices."),
    ("Threat Response Efficiency", "The automated 'Ack' function pipeline significantly reduces manual response time bounds (MTTR)."),
    ("Alert Accuracy", "The correlation scoring algorithm logs target records with minimal background false positives."),
    ("Alert Accuracy", "The system engine properly prioritizes critical infrastructure nodes on the network grid."),
    ("Overall Satisfaction", "Overall, the system delivers an optimal technical solution for local network threat profiling.")
]

sections_list = ["Ease of Use", "Alert Clarity", "Dashboard Readability", "Threat Response Efficiency", "Alert Accuracy", "Overall Satisfaction"]

q_global_num = 1
for current_section in sections_list:
    sec_qs = [q for q in questions if q[0] == current_section]
    if not sec_qs:
        continue
        
    add_section_banner(doc, current_section)
    
    # Generate clean table layout metrics matrix for the active section group
    grid = doc.add_table(rows=1, cols=6)
    grid.alignment = WD_TABLE_ALIGNMENT.CENTER
    grid.autofit = False
    apply_table_borders(grid, COLOR_BORDER_HEX)
    
    # Configure precise widths
    grid.columns[0].width = Inches(3.8) # Question string column width alignment
    for c in range(1, 6):
        grid.columns[c].width = Inches(0.64) # Option column widths
        
    # Table Grid Columns Header Label Titles Row
    hdr_cells = grid.rows[0].cells
    hdr_cells[0].text = "Evaluation Criteria Statements"
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[0].paragraphs[0].runs[0].font.size = Pt(9.5)
    hdr_cells[0].paragraphs[0].runs[0].font.color.rgb = COLOR_PRIMARY
    apply_cell_margins(hdr_cells[0], top=100, bottom=100, left=120, right=120)
    
    opt_labels = ["SD", "D", "N", "A", "SA"]
    for idx_lbl, lbl in enumerate(opt_labels):
        cell_target = hdr_cells[idx_lbl + 1]
        cell_target.text = lbl
        p_tgt = cell_target.paragraphs[0]
        p_tgt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_tgt.runs[0].font.bold = True
        p_tgt.runs[0].font.size = Pt(9.5)
        p_tgt.runs[0].font.color.rgb = COLOR_PRIMARY
        apply_cell_margins(cell_target, top=100, bottom=100, left=40, right=40)
        
    # Inject actual questions matrix data records rows
    for _, q_text in sec_qs:
        row_cells = grid.add_row().cells
        
        # Alternating background zebra stripes effect for quick readability tracking
        if q_global_num % 2 != 0:
            for cell in row_cells:
                apply_cell_shading(cell, COLOR_ZEBRA_HEX)
                
        # Write Question text
        apply_cell_margins(row_cells[0], top=120, bottom=120, left=120, right=120)
        p_q = row_cells[0].paragraphs[0]
        p_q.paragraph_format.space_before = Pt(2)
        p_q.paragraph_format.space_after = Pt(2)
        
        r_num = p_q.add_run(f"Q{q_global_num}. ")
        r_num.font.bold = True
        r_num.font.size = Pt(10)
        r_num.font.color.rgb = COLOR_ACCENT
        
        r_txt = p_q.add_run(q_text)
        r_txt.font.size = Pt(10)
        
        # Add checkbox brackets under columns 1-5 instead of raw loose numbers
        for c in range(1, 6):
            cell_box = row_cells[c]
            apply_cell_margins(cell_box, top=120, bottom=120, left=40, right=40)
            cell_box.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p_box = cell_box.paragraphs[0]
            p_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_box = p_box.add_run(f"[  ] {c}")
            r_box.font.size = Pt(9)
            r_box.font.color.rgb = COLOR_MUTED
            
        q_global_num += 1

# ─── OPEN-ENDED QUALITATIVE DATA SECTION ─────────────────────────────────────
add_section_banner(doc, "Qualitative Feedback & Engineering Recommendations")

open_queries = [
    "Identify the most significant feature or operational matrix of this risk profiling system:",
    "Propose structural UI revisions or architectural feature additions for subsequent sprints:",
    "Provide additional comments or direct observations regarding deployment stability metrics:"
]

for idx_oq, oq_text in enumerate(open_queries):
    p_oq = doc.add_paragraph()
    p_oq.paragraph_format.space_before = Pt(14)
    p_oq.paragraph_format.space_after = Pt(6)
    p_oq.paragraph_format.keep_with_next = True
    
    r_oq_num = p_oq.add_run(f"{idx_oq + 11}. ")
    r_oq_num.font.bold = True
    r_oq_num.font.color.rgb = COLOR_ACCENT
    
    r_oq_txt = p_oq.add_run(oq_text)
    r_oq_txt.font.bold = True
    
    # Modern styled response baseline text write grids lines box frame
    for line_idx in range(3):
        p_line = doc.add_paragraph()
        p_line.paragraph_format.space_before = Pt(0)
        p_line.paragraph_format.space_after = Pt(14 if line_idx == 2 else 10)
        r_l = p_line.add_run("─" * 94)
        r_l.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)  # Very light clean separator gray line

# End Closing Sign-off Frame Block
p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_end.paragraph_format.space_before = Pt(24)
r_end = p_end.add_run("Thank you for your valuable participation. Please return this completed form to Obed Kiplimo.\nKabarak University Final Year Capstone Project © 2026")
r_end.font.italic = True
r_end.font.size = Pt(9.5)
r_end.font.color.rgb = COLOR_MUTED

# ─── SAVE AND COMPILE OUTPUT ─────────────────────────────────────────────────
desktop_path = "/home/toro/Desktop/UAT_Evaluation_Package.docx"
doc.save(desktop_path)
print(f"\n🚀 Success! An eye-catching corporate-styled UAT Evaluation document is ready at:\n👉 {desktop_path}")